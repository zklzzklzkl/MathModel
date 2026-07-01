from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "source"
OUT = ROOT / "code" / "outputs"
OUT.mkdir(parents=True, exist_ok=True)


def extract_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for ti, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"index": ti, "rows": rows})
    return {
        "file": str(path),
        "paragraph_count": len(paragraphs),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def sheet_summary(xlsx: Path) -> dict:
    wb = load_workbook(xlsx, data_only=True, read_only=True)
    workbook_summary = {"file": str(xlsx), "sheets": []}
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        df = pd.read_excel(xlsx, sheet_name=sheet_name)
        df.columns = [str(c).strip() for c in df.columns]
        non_empty_rows = int(df.dropna(how="all").shape[0])
        duplicate_rows = int(df.duplicated().sum())
        missing = {str(k): int(v) for k, v in df.isna().sum().items()}
        numeric_cols = df.select_dtypes(include="number").columns.tolist()
        stats = {}
        if numeric_cols:
            desc = df[numeric_cols].describe().transpose()
            for col, row in desc.iterrows():
                stats[str(col)] = {
                    key: (None if pd.isna(value) else float(value))
                    for key, value in row.to_dict().items()
                }
        preview = df.head(8).where(pd.notna(df), None).to_dict(orient="records")
        csv_path = OUT / f"sheet_{sheet_name}.csv"
        df.to_csv(csv_path, index=False, encoding="utf-8-sig")
        workbook_summary["sheets"].append(
            {
                "name": sheet_name,
                "max_row": ws.max_row,
                "max_column": ws.max_column,
                "data_rows": non_empty_rows,
                "columns": df.columns.tolist(),
                "dtypes": {str(k): str(v) for k, v in df.dtypes.items()},
                "missing": missing,
                "duplicate_rows": duplicate_rows,
                "numeric_stats": stats,
                "preview": preview,
                "csv_export": str(csv_path),
            }
        )
    return workbook_summary


def main() -> None:
    docx_path = SOURCE / "C题(1).docx"
    xlsx_path = SOURCE / "附件(1).xlsx"
    result = {
        "problem_statement": extract_docx(docx_path),
        "workbook": sheet_summary(xlsx_path),
    }
    out_path = OUT / "input_extract.json"
    out_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(out_path)


if __name__ == "__main__":
    main()
